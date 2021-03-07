#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import xlrd
import os
import shutil
from docx import Document


def 新建目录(相对目录):
    if not os.path.exists(相对目录):
        os.makedirs(相对目录)


def 删除目录及目录下文件(相对目录):
    shutil.rmtree(相对目录)


def 文件名(相对目录):
    当前目录 = os.path.abspath('.')
    绝对目录 = 当前目录 + '\\' + 相对目录
    return os.listdir(绝对目录)


def 测试用例附件(文件相对地址):
    项目案例 = xlrd.open_workbook(文件相对地址)
    新建目录('测试用例附件')
    sheet1 = 项目案例.sheet_by_index(0)
    第1行 = sheet1.row_values(0)
    第1列 = sheet1.col_values(0)
    for i in range(1, sheet1.nrows):
        row_temp = sheet1.row_values(i)
        document = Document()
        for j in range(10):
            x = str(第1行[j]) + '：'
            document.add_heading(x, level=3)
            y = str(row_temp[j])
            document.add_paragraph(y)
            document.add_paragraph('')
        document.add_heading('执行截图：', level=3)
        document.add_paragraph('')
        第1列[i] = 第1列[i].replace('/', '')
        document.save("测试用例附件/{}.docx".format(第1列[i]))


def 主函数():
    新建目录('测试用例')
    文件名列表 = 文件名('测试用例')
    计数 = 1
    for 文件名字 in 文件名列表:
        print("\r进度：{}/{}".format(计数, len(文件名列表)), end='')
        文件相对地址 = '测试用例/' + 文件名字
        测试用例附件(文件相对地址)
        计数 += 1


if __name__ == '__main__':
    主函数()
    # 删除目录及目录下文件('测试用例')
    # 删除目录及目录下文件('测试用例附件')
